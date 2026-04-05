"""Build the complete thesis as a single DOCX following the official format template.

Format specs (from 毕业论文3_14.docx):
- A4 (11906×16838 DXA)
- Margins: top/bottom 1417, left/right 1701 DXA
- 摘要/ABSTRACT/章标题: 黑体/TNR 小二号(18pt) 加粗居中, 段前段后1行
- 节标题(##): 黑体 小四号(12pt) 加粗, 左对齐
- 小节标题(###): 黑体 小四号(12pt) 加粗, 左对齐
- 正文: 宋体+TNR 小四号(12pt), 1.5倍行距, 首行缩进两字符
- 关键词: 黑体小四号加粗(标签) + 宋体小四号(内容)
- 图注: 宋体五号(10.5pt) 居中
- 表格: 宋体五号(10.5pt)
- 参考文献: 宋体五号(10.5pt) 悬挂缩进
- 页脚: 小五号(9pt) TNR 居中页码
- 公式: 可编辑OMML, 居中, 右侧标号 (章-序号)
"""
from __future__ import annotations

import re
from pathlib import Path

import latex2mathml.converter
from lxml import etree
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Inches, Pt, Emu, RGBColor, Twips

REBUILD_ROOT = Path(__file__).resolve().parent
OUTPUT_DOCX = REBUILD_ROOT / "output" / "doc" / "full_thesis.docx"
FIGURES_DIR = REBUILD_ROOT / "figures"

# All chapter source files in order
CHAPTER_SOURCES = [
    REBUILD_ROOT / "docs" / "chapter1_draft_cn.md",   # 摘要+ABSTRACT+第1章+参考文献
    REBUILD_ROOT / "docs" / "chapter2_final_cn.md",    # 第2章
    REBUILD_ROOT / "docs" / "chapter3_draft_cn.md",    # 第3章
    REBUILD_ROOT / "docs" / "chapter4_draft_cn.md",    # 第4章
    REBUILD_ROOT / "docs" / "chapter5_draft_cn.md",    # 第5章
]


# ============================================================
# LaTeX equations per chapter (raw LaTeX from markdown)
# ============================================================
# These map the normalized equation text to the LaTeX for OMML rendering
EQUATION_LATEX_MAP = {
    # Chapter 2
    r"M = \frac{|N_{\Delta>0}-N_{\Delta<0}|}{N-1},":
        r"M = \frac{|N_{\Delta>0}-N_{\Delta<0}|}{N-1}",
    r"T = |\mathrm{corr}(x,t)|.":
        r"T = |\mathrm{corr}(x,t)|",
    r"S = 0.5M + 0.5T.":
        r"S = 0.5M + 0.5T",
    r"c_n = \frac{\sum_{a=1}^{n} f(a)}{\sqrt{\sum_{a=1}^{n} f(a)^2}}.":
        r"c_n = \frac{\sum_{a=1}^{n} f(a)}{\sqrt{\sum_{a=1}^{n} f(a)^2}}",
    r"HI(t) = \frac{1}{K}\sum_{k=1}^{K}\left|\frac{c_k(t)-\mu_k}{\sigma_k}\right|,":
        r"HI(t) = \frac{1}{K}\sum_{k=1}^{K}\left|\frac{c_k(t)-\mu_k}{\sigma_k}\right|",
    r"\theta = \overline{HI}_0 + 3\sigma_{HI_0},":
        r"\theta = \overline{HI}_0 + 3\sigma_{HI_0}",
    # Chapter 3
    r"P_t = \frac{L_t}{L}":
        r"P_t = \frac{L_t}{L}",
    r"s = \frac{1}{N}\sum_{i=1}^{N} s_i, \quad s_i = \begin{cases} 2^{\,e_i/5}, & e_i \le 0 \\ 2^{-e_i/20}, & e_i > 0 \end{cases}":
        r"s = \frac{1}{N}\sum_{i=1}^{N} s_i, \quad s_i = \begin{cases} 2^{e_i/5}, & e_i \le 0 \\ 2^{-e_i/20}, & e_i > 0 \end{cases}",
}

# Figure maps: h2 heading → list of (rel_path, caption, width)
FIGURE_MAP = {
    "2.1 数据集与研究对象": [
        ("chapter2/fig2_1_pipeline.png", "图2-1 第2章固定12特征健康状态识别总体流程", 6.0),
        ("chapter2/fig2_2_dataset_distribution.png", "图2-2 XJTU-SY数据集各工况轴承寿命分布", 6.0),
    ],
    "2.2 水平振动信号预处理": [
        ("chapter2/fig2_3_wavelet_denoise.png", "图2-3 代表样本局部时窗的水平振动信号去噪前后对比", 6.0),
    ],
    "2.3 候选特征提取与初步筛选": [
        ("chapter2/fig2_4_feature_examples.png", "图2-4 代表轴承的候选时域特征序列示例", 6.0),
    ],
    "2.4 特征评分与建模特征确定": [
        ("chapter2/fig2_5_feature_selection.png", "图2-5 13维候选特征评分排序结果", 5.8),
        ("chapter2/fig2_6_selected_correlation.png", "图2-6 固定12特征的平均Pearson相关矩阵", 5.8),
    ],
    "2.5 累积变换与多特征融合健康指标": [
        ("chapter2/fig2_7_health_indicator.png", "图2-7 代表轴承的健康指标演化曲线", 5.8),
    ],
    "3.5 基线对比结果与分析": [
        ("chapter3/fig3_1_rmse_comparison.png", "图3-1 三工况基线模型测试集RMSE对比", 5.5),
        ("chapter3/fig3_2_score_comparison.png", "图3-2 三工况基线模型测试集score s对比", 5.5),
        ("chapter3/fig3_3_loss_curves.png", "图3-3 基线模型训练过程损失曲线", 6.0),
        ("chapter3/fig3_4_rul_prediction.png", "图3-4 测试集轴承RUL预测曲线对比", 6.0),
    ],
    "4.3 Optuna 搜索结果": [
        ("chapter4/fig4_1_trial_history.png", "图4-1 Optuna搜索过程目标值变化曲线", 5.5),
    ],
    "4.4 自动调参正式确认结果": [
        ("chapter4/fig4_2_baseline_vs_optuna.png", "图4-2 手动调参与Optuna调参后Transformer测试集RMSE对比", 5.5),
        ("chapter4/fig4_3_loss_optuna.png", "图4-3 Optuna优化后Transformer训练损失曲线", 6.0),
        ("chapter4/fig4_4_rul_comparison.png", "图4-4 基线与Optuna Transformer测试集RUL预测对比", 6.0),
    ],
}


# ============================================================
# OMML Namespaces
# ============================================================
MATH_NS = "http://schemas.openxmlformats.org/officeDocument/2006/math"
WORD_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
MATHML_NS = "http://www.w3.org/1998/Math/MathML"


def _m(tag):
    """Create element in math namespace."""
    return OxmlElement(f"m:{tag}")


def _m_run(text, italic=False):
    """Create an OMML run: m:r > m:rPr + m:t."""
    r = _m("r")
    # Run properties
    rpr = _m("rPr")
    # Set font style
    sty = _m("sty")
    if italic:
        sty.set(qn("m:val"), "i")
    else:
        sty.set(qn("m:val"), "p")  # plain
    rpr.append(sty)
    r.append(rpr)
    # Text
    t = _m("t")
    t.text = text
    r.append(t)
    return r


def _mathml_to_omml_children(parent_elem, omml_parent):
    """Recursively convert MathML children to OMML elements."""
    ns = MATHML_NS
    children = list(parent_elem)

    i = 0
    while i < len(children):
        child = children[i]
        tag = etree.QName(child).localname

        if tag == "mrow":
            _convert_mrow(child, omml_parent)
        elif tag == "mi":
            txt = (child.text or "").strip()
            if txt:
                omml_parent.append(_m_run(txt, italic=True))
        elif tag == "mn":
            txt = (child.text or "").strip()
            if txt:
                omml_parent.append(_m_run(txt, italic=False))
        elif tag == "mo":
            txt = (child.text or "").strip()
            if txt:
                # Check if this is a fence delimiter
                fence = child.get("fence", "false")
                form = child.get("form", "")
                if fence == "true":
                    # Will be handled by parent mrow's fence detection
                    omml_parent.append(_m_run(txt, italic=False))
                else:
                    omml_parent.append(_m_run(txt, italic=False))
        elif tag == "mfrac":
            _convert_mfrac(child, omml_parent)
        elif tag == "msub":
            _convert_msub(child, omml_parent)
        elif tag == "msup":
            _convert_msup(child, omml_parent)
        elif tag == "msubsup":
            _convert_msubsup(child, omml_parent)
        elif tag == "msqrt":
            _convert_msqrt(child, omml_parent)
        elif tag == "mover":
            _convert_mover(child, omml_parent)
        elif tag == "munder":
            _convert_munder(child, omml_parent)
        elif tag == "mtext":
            txt = (child.text or "").strip()
            if txt:
                omml_parent.append(_m_run(txt, italic=False))
        else:
            # Unknown tag, try to process children
            _mathml_to_omml_children(child, omml_parent)

        i += 1


def _convert_mrow(elem, omml_parent):
    """Convert mrow - check for fence delimiters first."""
    ns = MATHML_NS
    children = list(elem)

    # Check if this mrow is a fenced group: starts with fence prefix, ends with fence postfix
    if len(children) >= 2:
        first = children[0]
        last = children[-1]
        first_tag = etree.QName(first).localname
        last_tag = etree.QName(last).localname

        if (first_tag == "mo" and first.get("fence") == "true" and first.get("form") == "prefix"
                and last_tag == "mo" and last.get("fence") == "true" and last.get("form") == "postfix"):
            # Create delimiter element
            d = _m("d")
            dpr = _m("dPr")
            beg = _m("begChr")
            beg.set(qn("m:val"), first.text or "(")
            end_chr = _m("endChr")
            end_chr.set(qn("m:val"), last.text or ")")
            dpr.append(beg)
            dpr.append(end_chr)
            d.append(dpr)
            e = _m("e")
            # Process inner children (skip first and last fence)
            for child in children[1:-1]:
                tag = etree.QName(child).localname
                if tag == "mrow":
                    _convert_mrow(child, e)
                elif tag == "mi":
                    txt = (child.text or "").strip()
                    if txt:
                        e.append(_m_run(txt, italic=True))
                elif tag == "mn":
                    txt = (child.text or "").strip()
                    if txt:
                        e.append(_m_run(txt, italic=False))
                elif tag == "mo":
                    txt = (child.text or "").strip()
                    if txt:
                        e.append(_m_run(txt, italic=False))
                elif tag == "mfrac":
                    _convert_mfrac(child, e)
                elif tag == "msub":
                    _convert_msub(child, e)
                elif tag == "msup":
                    _convert_msup(child, e)
                elif tag == "msubsup":
                    _convert_msubsup(child, e)
                elif tag == "msqrt":
                    _convert_msqrt(child, e)
                elif tag == "mover":
                    _convert_mover(child, e)
                else:
                    _mathml_to_omml_children(child, e)
            d.append(e)
            omml_parent.append(d)
            return

    # Normal mrow: just process children
    _mathml_to_omml_children(elem, omml_parent)


def _convert_mfrac(elem, omml_parent):
    """mfrac → m:f."""
    children = list(elem)
    f = _m("f")
    # numerator
    num = _m("num")
    if len(children) > 0:
        _mathml_to_omml_children(children[0], num) if etree.QName(children[0]).localname == "mrow" else _convert_single(children[0], num)
    f.append(num)
    # denominator
    den = _m("den")
    if len(children) > 1:
        _mathml_to_omml_children(children[1], den) if etree.QName(children[1]).localname == "mrow" else _convert_single(children[1], den)
    f.append(den)
    omml_parent.append(f)


def _convert_single(elem, omml_parent):
    """Convert a single MathML element into omml_parent."""
    tag = etree.QName(elem).localname
    if tag == "mrow":
        _convert_mrow(elem, omml_parent)
    elif tag == "mi":
        txt = (elem.text or "").strip()
        if txt:
            omml_parent.append(_m_run(txt, italic=True))
    elif tag == "mn":
        txt = (elem.text or "").strip()
        if txt:
            omml_parent.append(_m_run(txt, italic=False))
    elif tag == "mo":
        txt = (elem.text or "").strip()
        if txt:
            omml_parent.append(_m_run(txt, italic=False))
    elif tag == "mfrac":
        _convert_mfrac(elem, omml_parent)
    elif tag == "msub":
        _convert_msub(elem, omml_parent)
    elif tag == "msup":
        _convert_msup(elem, omml_parent)
    elif tag == "msubsup":
        _convert_msubsup(elem, omml_parent)
    elif tag == "msqrt":
        _convert_msqrt(elem, omml_parent)
    elif tag == "mover":
        _convert_mover(elem, omml_parent)
    elif tag == "munder":
        _convert_munder(elem, omml_parent)
    elif tag == "mtext":
        txt = (elem.text or "").strip()
        if txt:
            omml_parent.append(_m_run(txt, italic=False))
    else:
        _mathml_to_omml_children(elem, omml_parent)


def _convert_msub(elem, omml_parent):
    """msub → m:sSub."""
    children = list(elem)
    ssub = _m("sSub")
    e = _m("e")
    if len(children) > 0:
        _convert_single(children[0], e)
    ssub.append(e)
    sub = _m("sub")
    if len(children) > 1:
        _convert_single(children[1], sub)
    ssub.append(sub)
    omml_parent.append(ssub)


def _convert_msup(elem, omml_parent):
    """msup → m:sSup."""
    children = list(elem)
    ssup = _m("sSup")
    e = _m("e")
    if len(children) > 0:
        _convert_single(children[0], e)
    ssup.append(e)
    sup = _m("sup")
    if len(children) > 1:
        _convert_single(children[1], sup)
    ssup.append(sup)
    omml_parent.append(ssup)


def _convert_msubsup(elem, omml_parent):
    """msubsup → m:nary (for Σ/∏) or m:sSubSup."""
    children = list(elem)
    if len(children) < 3:
        _mathml_to_omml_children(elem, omml_parent)
        return

    base_elem = children[0]
    base_tag = etree.QName(base_elem).localname
    base_text = (base_elem.text or "").strip() if base_tag == "mo" else ""

    # Check if this is a summation/product n-ary operator
    nary_chars = {"∑": "∑", "\u2211": "∑", "∏": "∏", "\u220f": "∏"}
    if base_text in nary_chars:
        nary = _m("nary")
        narypr = _m("naryPr")
        chr_elem = _m("chr")
        chr_elem.set(qn("m:val"), nary_chars[base_text])
        narypr.append(chr_elem)
        limloc = _m("limLoc")
        limloc.set(qn("m:val"), "undOvr")
        narypr.append(limloc)
        nary.append(narypr)
        # sub
        sub = _m("sub")
        _convert_single(children[1], sub)
        nary.append(sub)
        # sup
        sup = _m("sup")
        _convert_single(children[2], sup)
        nary.append(sup)
        # empty body (content follows as siblings in MathML)
        e = _m("e")
        nary.append(e)
        omml_parent.append(nary)
    else:
        # Generic sub-superscript
        ssubsup = _m("sSubSup")
        e = _m("e")
        _convert_single(children[0], e)
        ssubsup.append(e)
        sub = _m("sub")
        _convert_single(children[1], sub)
        ssubsup.append(sub)
        sup = _m("sup")
        _convert_single(children[2], sup)
        ssubsup.append(sup)
        omml_parent.append(ssubsup)


def _convert_msqrt(elem, omml_parent):
    """msqrt → m:rad."""
    rad = _m("rad")
    radpr = _m("radPr")
    deghide = _m("degHide")
    deghide.set(qn("m:val"), "1")
    radpr.append(deghide)
    rad.append(radpr)
    deg = _m("deg")
    rad.append(deg)
    e = _m("e")
    _mathml_to_omml_children(elem, e)
    rad.append(e)
    omml_parent.append(rad)


def _convert_mover(elem, omml_parent):
    """mover → m:bar (overline) or m:acc."""
    children = list(elem)
    if len(children) < 2:
        _mathml_to_omml_children(elem, omml_parent)
        return

    accent_elem = children[1]
    accent_text = (accent_elem.text or "").strip()

    # Check for overline (horizontal bar)
    if accent_text in ("¯", "‾", "\u00af", "\u203e", "\u2015", "―"):
        bar = _m("bar")
        barpr = _m("barPr")
        pos = _m("pos")
        pos.set(qn("m:val"), "top")
        barpr.append(pos)
        bar.append(barpr)
        e = _m("e")
        _convert_single(children[0], e)
        bar.append(e)
        omml_parent.append(bar)
    else:
        # Generic accent
        acc = _m("acc")
        accpr = _m("accPr")
        chr_elem = _m("chr")
        chr_elem.set(qn("m:val"), accent_text)
        accpr.append(chr_elem)
        acc.append(accpr)
        e = _m("e")
        _convert_single(children[0], e)
        acc.append(e)
        omml_parent.append(acc)


def _convert_munder(elem, omml_parent):
    """munder → render base with subscript-like underscript."""
    children = list(elem)
    if len(children) < 2:
        _mathml_to_omml_children(elem, omml_parent)
        return
    # Simple approach: render as subscript
    ssub = _m("sSub")
    e = _m("e")
    _convert_single(children[0], e)
    ssub.append(e)
    sub = _m("sub")
    _convert_single(children[1], sub)
    ssub.append(sub)
    omml_parent.append(ssub)


def latex_to_omml(latex_str: str):
    """Convert LaTeX string to OMML oMath element."""
    mathml_str = latex2mathml.converter.convert(latex_str)
    tree = etree.fromstring(mathml_str.encode())

    omath = _m("oMath")
    # The root is <math>, process its children
    _mathml_to_omml_children(tree, omath)
    return omath


# ============================================================
# Font / Style
# ============================================================

def set_run_font(run, cn_font="宋体", en_font="Times New Roman", size_pt=12, bold=False):
    run.font.name = en_font
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:eastAsia"), cn_font)
    rfonts.set(qn("w:ascii"), en_font)
    rfonts.set(qn("w:hAnsi"), en_font)
    run.font.size = Pt(size_pt)
    run.bold = bold


def configure_styles(doc: Document):
    """Set up styles per the template."""
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    rpr = normal._element.get_or_add_rPr()
    rf = OxmlElement("w:rFonts")
    rf.set(qn("w:eastAsia"), "宋体")
    rf.set(qn("w:ascii"), "Times New Roman")
    rf.set(qn("w:hAnsi"), "Times New Roman")
    rpr.insert(0, rf)
    normal.font.size = Pt(12)  # 小四号
    pf = normal.paragraph_format
    pf.line_spacing = 1.5
    pf.first_line_indent = Pt(24)  # 两字符
    pf.space_after = Pt(0)
    pf.space_before = Pt(0)

    # Heading 1: 章标题 - 小二号(18pt) 黑体 加粗 居中
    h1 = doc.styles["Heading 1"]
    h1.font.name = "Times New Roman"
    h1_rpr = h1._element.get_or_add_rPr()
    h1_rf = OxmlElement("w:rFonts")
    h1_rf.set(qn("w:eastAsia"), "黑体")
    h1_rf.set(qn("w:ascii"), "Times New Roman")
    h1_rf.set(qn("w:hAnsi"), "Times New Roman")
    h1_rpr.insert(0, h1_rf)
    h1.font.size = Pt(18)
    h1.font.bold = True
    h1.font.color.rgb = RGBColor(0, 0, 0)
    h1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    h1.paragraph_format.space_before = Pt(12)
    h1.paragraph_format.space_after = Pt(12)
    h1.paragraph_format.first_line_indent = Pt(0)
    h1.paragraph_format.line_spacing = 1.5
    h1.paragraph_format.page_break_before = True

    # Heading 2: 节标题 - 小四号(12pt) 黑体 加粗 左对齐
    h2 = doc.styles["Heading 2"]
    h2.font.name = "Times New Roman"
    h2_rpr = h2._element.get_or_add_rPr()
    h2_rf = OxmlElement("w:rFonts")
    h2_rf.set(qn("w:eastAsia"), "黑体")
    h2_rf.set(qn("w:ascii"), "Times New Roman")
    h2_rf.set(qn("w:hAnsi"), "Times New Roman")
    h2_rpr.insert(0, h2_rf)
    h2.font.size = Pt(12)
    h2.font.bold = True
    h2.font.color.rgb = RGBColor(0, 0, 0)
    h2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    h2.paragraph_format.space_before = Pt(6)
    h2.paragraph_format.space_after = Pt(4)
    h2.paragraph_format.first_line_indent = Pt(0)
    h2.paragraph_format.line_spacing = 1.5

    # Heading 3: 小节标题 - 小四号(12pt) 黑体 加粗 左对齐
    h3 = doc.styles["Heading 3"]
    h3.font.name = "Times New Roman"
    h3_rpr = h3._element.get_or_add_rPr()
    h3_rf = OxmlElement("w:rFonts")
    h3_rf.set(qn("w:eastAsia"), "黑体")
    h3_rf.set(qn("w:ascii"), "Times New Roman")
    h3_rf.set(qn("w:hAnsi"), "Times New Roman")
    h3_rpr.insert(0, h3_rf)
    h3.font.size = Pt(12)
    h3.font.bold = True
    h3.font.color.rgb = RGBColor(0, 0, 0)
    h3.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    h3.paragraph_format.space_before = Pt(4)
    h3.paragraph_format.space_after = Pt(2)
    h3.paragraph_format.first_line_indent = Pt(0)
    h3.paragraph_format.line_spacing = 1.5


def setup_page(doc: Document):
    section = doc.sections[0]
    section.page_width = Emu(11906 * 635)
    section.page_height = Emu(16838 * 635)
    section.top_margin = Emu(1417 * 635)
    section.bottom_margin = Emu(1417 * 635)
    section.left_margin = Emu(1701 * 635)
    section.right_margin = Emu(1701 * 635)


# ============================================================
# Text helpers
# ============================================================

def format_text(text: str) -> str:
    def replace_inline_math(match):
        expr = match.group(1)
        for src, dst in {
            r"\{": "{", r"\}": "}", r"\dots": "...",
            r"\Delta": "\u0394", r"\mu": "\u03bc", r"\sigma": "\u03c3",
            r"\theta": "\u03b8", r"\mathrm{corr}": "corr",
            r"\overline{HI}": "mean(HI)",
        }.items():
            expr = expr.replace(src, dst)
        expr = expr.replace("{", "(").replace("}", ")").replace("\\", "")
        return expr
    text = re.sub(r"\$(.+?)\$", replace_inline_math, text)
    return text.replace("`", "")


# ============================================================
# Content builders
# ============================================================

def add_body_paragraph(doc, text, is_english=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    cn_font = "宋体"
    en_font = "Times New Roman"
    # Handle **bold** inline
    parts = re.split(r"(\*\*.+?\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = p.add_run(format_text(part[2:-2]))
            set_run_font(run, cn_font if not is_english else en_font, en_font, 12, bold=True)
        elif part:
            run = p.add_run(format_text(part))
            set_run_font(run, cn_font if not is_english else en_font, en_font, 12)
    return p


def add_chapter_title(doc, text):
    """Add H1 chapter title via heading style (gets page break, centered, 18pt bold)."""
    p = doc.add_heading(text, level=1)
    for run in p.runs:
        set_run_font(run, "黑体", "Times New Roman", 18, bold=True)
    return p


def add_abstract_title(doc, text, is_english=False):
    """摘要/ABSTRACT title: 小二号加粗居中, with page break."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Pt(0)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(12)
    p.paragraph_format.page_break_before = True
    run = p.add_run(text)
    if is_english:
        set_run_font(run, "Times New Roman", "Times New Roman", 18, bold=True)
    else:
        set_run_font(run, "黑体", "Times New Roman", 18, bold=True)
    return p


def add_keywords(doc, text, is_english=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Pt(0)
    p.paragraph_format.space_before = Pt(6)
    sep = "：" if "：" in text else (":" if ":" in text else "")
    if sep:
        label, content = text.split(sep, 1)
    else:
        label, content = text, ""
    label = label.replace("**", "").strip()
    content = content.replace("**", "").strip()
    cn = "宋体" if not is_english else "Times New Roman"
    en = "Times New Roman"
    run = p.add_run(label + sep + " ")
    set_run_font(run, "黑体" if not is_english else en, en, 12, bold=True)
    if content:
        run2 = p.add_run(content)
        set_run_font(run2, cn, en, 12)


def add_equation_omml(doc, raw_text, chapter_num, eq_counter):
    """Add editable OMML equation with right-aligned numbering (chapter-counter).

    Returns the updated equation counter.
    """
    normalized = " ".join(raw_text.split())
    latex_str = EQUATION_LATEX_MAP.get(normalized)

    if latex_str is None:
        # Fallback: render as plain text
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.first_line_indent = Pt(0)
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run(normalized)
        set_run_font(run, "Times New Roman", "Times New Roman", 11)
        return eq_counter

    eq_counter += 1
    eq_label = f"({chapter_num}-{eq_counter})"

    # Create paragraph
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT  # we'll use tabs for centering + right align
    p.paragraph_format.first_line_indent = Pt(0)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)

    # Set up tab stops: center tab at middle, right tab at right margin
    # Page width ~6.3in usable (A4 with margins), center ≈ 3.15in, right ≈ 6.3in
    ppr = p._element.get_or_add_pPr()
    tabs = OxmlElement("w:tabs")
    # Center tab for equation
    tab_center = OxmlElement("w:tab")
    tab_center.set(qn("w:val"), "center")
    tab_center.set(qn("w:pos"), "4680")  # ~3.25 inches in twips
    tabs.append(tab_center)
    # Right tab for equation number
    tab_right = OxmlElement("w:tab")
    tab_right.set(qn("w:val"), "right")
    tab_right.set(qn("w:pos"), "9360")  # ~6.5 inches in twips
    tabs.append(tab_right)
    ppr.append(tabs)

    # Tab to center position
    run_tab1 = p.add_run()
    run_tab1._element.append(OxmlElement("w:tab"))

    # Insert OMML equation
    try:
        omath = latex_to_omml(latex_str)
        p._element.append(omath)
    except Exception as e:
        print(f"  OMML conversion failed for: {latex_str[:50]}... Error: {e}")
        run = p.add_run(normalized)
        set_run_font(run, "Times New Roman", "Times New Roman", 11)

    # Tab to right position, then add equation number
    run_tab2 = p.add_run()
    run_tab2._element.append(OxmlElement("w:tab"))
    run_num = p.add_run(eq_label)
    set_run_font(run_num, "Times New Roman", "Times New Roman", 12)

    return eq_counter


def add_figure(doc, rel_path, caption, width_inches=6.0):
    img = FIGURES_DIR / rel_path
    if not img.exists():
        print(f"  WARNING: figure not found: {img}")
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = Pt(0)
    run = p.add_run()
    run.add_picture(str(img), width=Inches(width_inches))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.first_line_indent = Pt(0)
    cap.paragraph_format.space_after = Pt(6)
    run = cap.add_run(caption)
    set_run_font(run, "宋体", "Times New Roman", 10.5)


def _add_cell_with_math(p, text, cn_font="宋体", en_font="Times New Roman",
                        size_pt=10.5, bold=False):
    """Add cell content with inline $...$ rendered as OMML math."""
    parts = re.split(r"(\$[^$]+\$)", text)
    for part in parts:
        if part.startswith("$") and part.endswith("$"):
            latex_str = part[1:-1]
            try:
                omath = latex_to_omml(latex_str)
                p._element.append(omath)
            except Exception:
                run = p.add_run(format_text(part))
                set_run_font(run, cn_font, en_font, size_pt, bold=bold)
        elif part:
            run = p.add_run(format_text(part))
            set_run_font(run, cn_font, en_font, size_pt, bold=bold)


def add_table_from_markdown(doc, md_lines):
    rows_data = []
    for line in md_lines:
        line = line.strip()
        if not line.startswith("|"):
            continue
        cells = [c.strip() for c in line.split("|")[1:-1]]
        if all(set(c) <= set("- :") for c in cells):
            continue
        rows_data.append(cells)
    if not rows_data:
        return
    num_cols = len(rows_data[0])
    table = doc.add_table(rows=len(rows_data), cols=num_cols)
    table.style = "Table Grid"
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for ri, row_cells in enumerate(rows_data):
        for ci, cell_text in enumerate(row_cells):
            if ci >= num_cols:
                break
            cell = table.rows[ri].cells[ci]
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.first_line_indent = Pt(0)
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            has_math = "$" in cell_text
            if ri == 0:
                if has_math:
                    _add_cell_with_math(p, cell_text, "黑体", "Times New Roman", 10.5, bold=True)
                else:
                    run = p.add_run(cell_text)
                    set_run_font(run, "黑体", "Times New Roman", 10.5, bold=True)
            else:
                if has_math:
                    _add_cell_with_math(p, cell_text, "宋体", "Times New Roman", 10.5)
                else:
                    run = p.add_run(cell_text)
                    set_run_font(run, "宋体", "Times New Roman", 10.5)


def add_numbered_item(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Pt(0)
    p.paragraph_format.left_indent = Pt(24)
    parts = re.split(r"(\*\*.+?\*\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = p.add_run(format_text(part[2:-2]))
            set_run_font(run, "宋体", "Times New Roman", 12, bold=True)
        elif part:
            run = p.add_run(format_text(part))
            set_run_font(run, "宋体", "Times New Roman", 12)


def add_reference_item(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.left_indent = Pt(21)
    p.paragraph_format.first_line_indent = Pt(-21)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(format_text(text))
    set_run_font(run, "宋体", "Times New Roman", 10.5)


def insert_figures_for_h2(doc, h2_heading):
    """Insert figures associated with the given h2 heading."""
    figures = FIGURE_MAP.get(h2_heading, [])
    for rel_path, caption, width in figures:
        add_figure(doc, rel_path, caption, width)


# ============================================================
# Markdown parser
# ============================================================

def parse_markdown(path: Path):
    text = path.read_text(encoding="utf-8")
    blocks = []
    para_lines = []
    table_lines = []
    in_math = False
    math_lines = []

    def flush_para():
        nonlocal para_lines
        if para_lines:
            blocks.append(("paragraph", " ".join(l.strip() for l in para_lines).strip()))
        para_lines = []

    def flush_math():
        nonlocal math_lines
        if math_lines:
            blocks.append(("equation", " ".join(l.strip() for l in math_lines).strip()))
        math_lines = []

    def flush_table():
        nonlocal table_lines
        if table_lines:
            blocks.append(("table", list(table_lines)))
        table_lines = []

    for raw in text.splitlines():
        line = raw.rstrip()
        if line.strip() == "$$":
            flush_para(); flush_table()
            if in_math:
                flush_math()
            in_math = not in_math
            continue
        if in_math:
            math_lines.append(line); continue
        if line.strip().startswith("|"):
            flush_para(); table_lines.append(line); continue
        elif table_lines:
            flush_table()
        if line.startswith("# "):
            flush_para(); blocks.append(("h1", line[2:].strip())); continue
        if line.startswith("## "):
            flush_para(); blocks.append(("h2", line[3:].strip())); continue
        if line.startswith("### "):
            flush_para(); blocks.append(("h3", line[4:].strip())); continue
        if not line.strip():
            flush_para(); continue
        if re.match(r"^\d+\.\s+", line.strip()):
            flush_para(); blocks.append(("numbered", line.strip())); continue
        para_lines.append(line)
    flush_para(); flush_math(); flush_table()
    return blocks


# ============================================================
# Main builder
# ============================================================

def build():
    OUTPUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    configure_styles(doc)
    setup_page(doc)

    # State
    in_abstract_cn = False
    in_abstract_en = False
    in_references = False
    skip_section = False  # skip "参考文献说明" etc.
    is_first_title = True  # suppress page break for very first element
    current_chapter_num = 0
    eq_counter = 0  # per-chapter equation counter
    current_h2 = ""

    # Deferred reference blocks (to be rendered at the very end)
    deferred_references = []

    for source_path in CHAPTER_SOURCES:
        blocks = parse_markdown(source_path)
        current_h2 = ""
        in_references = False
        skip_section = False

        for idx, (block_type, content) in enumerate(blocks):

            # ── Look ahead: if the next block is a new h2 or h1, flush figures for current h2 ──
            def maybe_flush_figures():
                nonlocal current_h2
                if current_h2:
                    insert_figures_for_h2(doc, current_h2)

            # ── H1 titles ──
            if block_type == "h1":
                # Flush figures from previous h2 before starting new section
                maybe_flush_figures()
                current_h2 = ""
                skip_section = False

                if "摘" in content and "要" in content:
                    in_abstract_cn = True
                    in_abstract_en = False
                    in_references = False
                    p = add_abstract_title(doc, "摘  要", is_english=False)
                    if is_first_title:
                        p.paragraph_format.page_break_before = False
                        is_first_title = False
                    continue
                elif content.strip() == "ABSTRACT":
                    in_abstract_cn = False
                    in_abstract_en = True
                    in_references = False
                    add_abstract_title(doc, "ABSTRACT", is_english=True)
                    is_first_title = False
                    continue
                else:
                    in_abstract_cn = False
                    in_abstract_en = False
                    in_references = False
                    # Extract chapter number
                    ch_match = re.search(r"第(\d)章", content)
                    if ch_match:
                        current_chapter_num = int(ch_match.group(1))
                        eq_counter = 0  # reset per chapter
                    p = add_chapter_title(doc, content)
                    if is_first_title:
                        p.paragraph_format.page_break_before = False
                        is_first_title = False
                    continue

            # ── Defer reference blocks ──
            if in_references:
                deferred_references.append((block_type, content))
                continue

            # ── Skip section blocks ──
            if skip_section:
                if block_type in ("h1", "h2"):
                    skip_section = False
                    # fall through to h2 handler below
                else:
                    continue

            # ── H2 / H3 ──
            if block_type == "h2":
                # Flush figures from previous h2
                if current_h2:
                    insert_figures_for_h2(doc, current_h2)

                in_abstract_cn = False
                in_abstract_en = False

                # Check for references section (exact match) → defer to end
                if content.strip() == "参考文献":
                    in_references = True
                    current_h2 = ""
                    continue

                # Check for "参考文献说明" → skip entirely
                if "参考文献说明" in content:
                    skip_section = True
                    current_h2 = ""
                    continue

                current_h2 = content
                h = doc.add_heading(content, level=2)
                for run in h.runs:
                    set_run_font(run, "黑体", "Times New Roman", 12, bold=True)
                continue

            if block_type == "h3":
                h = doc.add_heading(content, level=3)
                for run in h.runs:
                    set_run_font(run, "黑体", "Times New Roman", 12, bold=True)
                continue

            # ── Keywords ──
            if block_type == "paragraph" and ("关键词" in content or "Key Words" in content):
                add_keywords(doc, content, is_english=in_abstract_en)
                continue

            # ── Normal content ──
            if block_type == "paragraph":
                add_body_paragraph(doc, content, is_english=in_abstract_en)
            elif block_type == "equation":
                eq_counter = add_equation_omml(doc, content, current_chapter_num, eq_counter)
            elif block_type == "table":
                add_table_from_markdown(doc, content)
            elif block_type == "numbered":
                add_numbered_item(doc, content)

        # End of chapter source file: flush figures for last h2
        if current_h2:
            insert_figures_for_h2(doc, current_h2)
            current_h2 = ""

    # ── Render deferred references at the very end ──
    if deferred_references:
        # Add "参考文献" as H1 heading
        p = add_chapter_title(doc, "参考文献")
        for block_type, content in deferred_references:
            if block_type == "paragraph" and re.search(r"\[\d+\]", content):
                add_reference_item(doc, content)
            elif block_type == "paragraph":
                add_body_paragraph(doc, content)
            elif block_type == "numbered":
                add_numbered_item(doc, content)

    doc.save(OUTPUT_DOCX)
    print(f"Saved: {OUTPUT_DOCX}")
    return OUTPUT_DOCX


if __name__ == "__main__":
    build()
